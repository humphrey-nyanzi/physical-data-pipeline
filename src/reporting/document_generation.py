"""
Document Generation Module for DOCX Report Creation

Low-level utilities for creating formatted Word documents from GPS
performance analysis data. Handles styling, table insertion, visualization
embedding, and document structure.

Abstracts away python-docx API details to provide high-level document
building functions for the reporting pipeline.

Author: Performance Analytics Team
Version: 1.0
"""

import os
import math
from io import BytesIO
from docx.enum.text import WD_ALIGN_PARAGRAPH
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.section import WD_ORIENT, WD_SECTION
from src.config.speed_zones import get_speed_zones


from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def add_table_of_contents(doc: Document) -> None:
    """
    Insert a Table of Contents field into the document.
    
    Forces Word to prompt for an update upon opening the document 
    to populate the TOC with actual page numbers and headings.
    """
    # 1. Add 'updateFields' setting to the document to trigger prompt on open
    element = doc.settings.element
    update_fields = element.find(qn('w:updateFields'))
    if update_fields is None:
        update_fields = OxmlElement('w:updateFields')
        update_fields.set(qn('w:val'), 'true')
        element.append(update_fields)

    # 2. Add the TOC Heading
    doc.add_heading("Table of Contents", level=1)
    
    # 3. Add the TOC Field
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()
    
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    fldChar1.set(qn('w:dirty'), 'true') # Mark as dirty to suggest update
    
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # \o "1-3": Use headings level 1 to 3
    # \h: Hyperlinks
    # \z: Hide tab leader and page numbers in Web view
    # \u: Use outline levels
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
    
    doc.add_page_break()


def add_table_caption(doc: Document, caption: str) -> None:
    """
    Add a formatted caption below or above a table with automatic numbering.
    """
    if not hasattr(doc, '_table_count'):
        doc._table_count = 0
    doc._table_count += 1
    
    full_caption = f"Table {doc._table_count}: {caption}"
    
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = paragraph.add_run(full_caption)
    run.italic = True
    # Make caption slightly smaller than standard text
    run.font.size = doc.styles['Normal'].font.size

def fmt_cell_value(val, float_fmt: str = "{:,.2f}") -> str:
    """
    Format a cell value safely for DOCX insertion.

    Handles None, NaN, integers, and floats appropriately.

    Args:
        val: Value to format
        float_fmt: Format string for floats (default: "{:,.2f}")

    Returns:
        Formatted string representation
    """
    if val is None or (isinstance(val, float) and math.isnan(val)):
        return ""
    if isinstance(val, int):
        return f"{val:,}"
    if isinstance(val, float):
        return float_fmt.format(val)
    return str(val)


def add_dataframe_as_table(
    doc: Document, df: pd.DataFrame, style: str = "List Table 2 - Accent 5", caption: str = None
) -> None:
    """
    Add a DataFrame to the document as a formatted table.

    Creates table with header row and data rows from DataFrame.
    Uses 'List Table 2 - Accent 5' style by default and removes cell spacing.

    Args:
        doc: python-docx Document object
        df: DataFrame to insert
        style: Table style name
        caption: Optional caption to display below the table
    """
    # Pre-allocate the entire table (header + data rows)
    # docx is much faster when rows are allocated at once rather than using add_row()
    num_rows = len(df) + 1
    num_cols = len(df.columns)
    table = doc.add_table(rows=num_rows, cols=num_cols)
    
    try:
        table.style = style
    except Exception:
        table.style = "Table Grid" # Fallback if style not found
        
    table.autofit = True
    
    # Autofit to window (100% width)
    if table._tbl.tblPr is not None and table._tbl.tblPr.xpath('w:tblW'):
        tblW = table._tbl.tblPr.xpath('w:tblW')[0]
        tblW.set(qn('w:type'), 'pct')
        tblW.set(qn('w:w'), '5000')

    # Add header row
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
        
        for paragraph in hdr_cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.space_before = Pt(0)
            paragraph.paragraph_format.space_after = Pt(0)
            for run in paragraph.runs:
                run.bold = True

    # Add data rows using pre-allocated table.rows
    for r_idx, (idx, row) in enumerate(df.iterrows(), start=1):
        row_cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            text = fmt_cell_value(val)
            row_cells[c_idx].text = text
            
            # Remove spacing and align
            col_name = str(df.columns[c_idx]).lower()
            for paragraph in row_cells[c_idx].paragraphs:
                paragraph.paragraph_format.space_before = Pt(0)
                paragraph.paragraph_format.space_after = Pt(0)
                if col_name in ['s/n', 'rank', 's/no', 'match day', 'value'] or 'count' in col_name:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
    if caption:
        add_table_caption(doc, caption)


def set_landscape(doc: Document) -> None:
    """Add a new landscape section to the document."""
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_width, new_height = new_section.page_height, new_section.page_width
    new_section.orientation = WD_ORIENT.LANDSCAPE
    new_section.page_width = new_width
    new_section.page_height = new_height


def set_portrait(doc: Document) -> None:
    """Add a new portrait section to the document."""
    new_section = doc.add_section(WD_SECTION.NEW_PAGE)
    new_section.orientation = WD_ORIENT.PORTRAIT
    new_section.page_width = Inches(8.5)
    new_section.page_height = Inches(11)


def add_branded_header(doc: Document, text: str, icon_path: str = None) -> None:
    """Add a heading with an optional icon for
     branding."""
    p = doc.add_heading('', level=1)
    if icon_path and os.path.exists(icon_path):
        run = p.add_run()
        try:
            run.add_picture(icon_path, width=Inches(0.35))
            run.add_text("  ") # Spacer
        except Exception as e:
            print(f"Warning: Could not add icon {icon_path}: {e}")
    p.add_run(text)


def add_title_page(doc: Document, title: str, subtitle: str, logo_path: str = None) -> None:
    """Create a  title page with logo and branding."""
    # Move to first page
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    if logo_path and os.path.exists(logo_path):
        run = p.add_run()
        try:
            run.add_picture(logo_path, width=Inches(2.5))
            doc.add_paragraph()
            doc.add_paragraph()
        except Exception as e:
            print(f"Warning: Could not add logo {logo_path}: {e}")
        
    t = doc.add_paragraph(title)
    t.style = 'Title'
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    st = doc.add_paragraph(subtitle)
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in st.runs:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(80, 80, 80)
    
    doc.add_page_break()

def add_introduction_section(doc: Document, club_name: str, season: str = "2025/26") -> None:
    """
    Add standardized Introduction section to report.

    Contains mission statement, GPS investment rationale,
    report objectives and organization.

    Args:
        doc: python-docx Document object
        club_name: Name of club for personalization
        season: Football season (e.g., '2025/26')
    """
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        """In a strategic move to elevate domestic football standards, the football federation invested substantially in GPS tracking systems for clubs across the top-tier leagues. By equipping teams with state-of-the-art performance monitoring technology, the federation demonstrates its commitment to player development, optimized match preparation, and sustained competitive excellence.

This report focuses on {}'s physical performance during the {} season, leveraging GPS tracking data to deliver actionable insights. The major objectives are to:
• Quantify player workloads and intensity metrics
• Identify performance trends and outliers
• Inform training load management, injury prevention, and tactical planning

It is important to note that accurate analysis and subsequent interpretation depend on consistent device usage, precise session tagging, and timely data synchronization. Where gaps arise due to technical or logistical constraints, we highlight their impact on the analyses.

The report is organized into seven sections:
• Methodology: Data collection, cleaning, and analysis protocols
• Key Concepts and Definitions in Physical Performance Analysis
• Season Results: Individual and aggregated performance metrics (e.g., total distance, high-intensity efforts, accelerations)
• Club Comparison: {} versus league averages to reveal relative strengths and areas for improvement
• Challenges: Recurring issues such as inconsistent device usage, mislabelled sessions, failure to upload, mismanagement of equipment
• Recommendations: Corrective measures including refresher training for GPS operators, compliance to session naming protocols
• Future Plans: Strategic initiatives to enhance data integration, reporting consistency, and long-term performance monitoring across all top-tier clubs
""".format(club_name, season, club_name)
    )


def add_methodology_section(doc: Document, season: str = "2025/26") -> None:
    """
    Add standardized Methodology section to report.

    Describes data collection protocols, cleaning procedures, and analysis approach.

    Args:
        doc: python-docx Document object
        season: Football season (e.g., '2025/26')
    """
    doc.add_heading("Methodology", level=1)
    doc.add_paragraph("""
This report is based on data collected using a GPS player tracking system during the {} football season. The data includes performance metrics such as total distance covered, sprint distance, work ratio, player load, session availability, top speed, accelerations, decelerations, energy, power plays, impacts, total actions, distance per minute.
""".format(season) + """
Data was collected and uploaded by trained club staff immediately after match sessions. The performance analytics unit oversaw the data collection process, ensuring sessions were split by halves, correctly labelled, and uploaded within seventy-two (72) hours post-match. Where necessary, additional follow-up was done with the club operators to correct mislabelled sessions or fill in missing data, ensuring full coverage of match data.

The rigorous data cleaning process involved removing duplicates, resolving inconsistencies in player-pod assignments, and verifying that each match session met the minimum completeness threshold (e.g., full ninety (90) minutes, sufficient number of tracked players). Any session failing these criteria was excluded from further analysis to maintain consistency.

Once cleaned, data was aggregated by individual player and by club. Metrics were normalized to account for variations in match frequency, squad size, and overall data availability, and clubs with fewer than ten complete sessions were omitted from comparative analyses. All statistical summaries and visualizations were generated using Python, following established workflows to ensure reliable, consistent benchmarks for technical review and long-term planning.
    """)

def add_key_concepts_section(doc: Document, season: str = "2025/26") -> None:
    """
    Add standardized Key Concepts and Definitions section to report.

    Provides educational content on physical performance metrics.

    Args:
        doc: python-docx Document object
        season: Football season (e.g., '2025/26')
    """
    doc.add_heading(
        "Key Concepts and Definitions in Physical Performance Analysis", level=1
    )

    doc.add_paragraph(
        """This section gives a basic introduction to coaches, analysts and any other personnel who may be interested in physical performance data. It includes basic definitions and brief descriptions of the metrics and terminology used in the field. It also includes a section of why a club should be interested in their physical data.

Physical performance analysis: A method of collecting and interpreting data on a player's movement and workload during training or matches. By measuring total and sprint distance, player load and work ratios, performance staff can monitor fitness adaptations, manage fatigue and implement injury-prevention protocols. Clubs can use these insights to adjust training loads and optimise readiness for crucial fixtures.

GPS tracking technology: A wearable athlete-monitoring system that integrates global positioning sensors with inertial measurement units (accelerometers, gyroscopes and magnetometers) to capture real-time movement data at high sampling rates. After each session, devices sync to base stations and upload into the tracking software suite, which calculates metrics such as sprint profiles, accelerations and overall player load. Clubs use this information to tailor individual conditioning programs and track recovery between matchdays.

Volume: This refers to the total amount of physical work performed by a player over a training session or match. It captures aggregated outputs such as total distance covered or total number of accelerations and decelerations. Volume metrics help coaches gauge overall workload across a full session. Volume metrics include total distance, sprint distance, power plays, energy, impacts, accelerations, decelerations, total actions and player load.

Intensity: This refers to the rate or magnitude at which physical work is carried out. It measures how quickly or forcefully a player executes movements, for example distance per minute or accelerations per minute. Clubs use intensity metrics to adjust training drills for peak performance on matchdays. Intensity metrics include top speed, distance per minute, work ratio, power score, acceleration counts per minute and deceleration counts per minute.
    """
    )

    doc.add_paragraph("GPS Metrics and Their Meaning")

    # Add metrics table
    metrics_data = {
        "Metric": [
            "Total Distance",
            "Sprint Distance",
            "Power Plays",
            "Energy",
            "Impacts",
            "Accelerations",
            "Decelerations",
            "Total Actions",
            "Player Load",
            "Top Speed",
            "Distance per Minute",
            "Work Ratio",
            "Power Score",
            "Acceleration Counts Per Minute",
            "Deceleration Counts Per Minute",
        ],
        "Definition": [
            "Total kilometres a player covers while walking or running in a session",
            "Distance covered while running above a set speed threshold of 25.2 km/h",
            "Number of intense actions a player was involved in",
            "Measure of how much energy is expended during a session",
            "Number of high-force events (e.g., collisions, landings, tackles)",
            "Total count of acceleration events above a set threshold",
            "Total count of deceleration events above a set threshold",
            "Sum of total accelerations and total decelerations",
            "Measures the overall workload on a player's body in a session",
            "Maximum speed a player reaches in a session",
            "Intensity measure showing metres covered every minute",
            "Ratio of high-intensity work time to total time",
            "Measures the power output used per kilogram of a player's weight",
            "Total accelerations divided by minutes played",
            "Total decelerations divided by minutes played",
        ],
        "Unit": [
            "km",
            "m",
            "Count",
            "kcal",
            "Count",
            "Count",
            "Count",
            "Count",
            "AU",
            "km/h",
            "m/min",
            "%",
            "w/kg",
            "counts/min",
            "counts/min",
        ],
        "Type": ["Volume"] * 9 + ["Intensity"] * 6,
    }

    metrics_df = pd.DataFrame(metrics_data)
    add_dataframe_as_table(doc, metrics_df)

    doc.add_paragraph()
    doc.add_paragraph(
        """Why should a club be interested in these metrics?

1. Monitoring these metrics allows coaches to identify fatigue or underperformance before it leads to injury or decline in form. For example, a sudden drop in total actions or distance per minute can signal that a player is not recovering adequately between sessions.

2. These measures enable position-specific benchmarks, ensuring defenders, wingers, midfielders and strikers train to the demands of their roles rather than a generic team standard. Comparing work ratio and power score across similar positions helps staff tailor workloads and maintain optimal performance levels.

3. Individualized training programs are designed using these data points. If a player's accelerations per minute fall below the team standard, practice drills can be adjusted to include more high-intensity runs and recovery phases, thereby bridging the gap to match demands.

4. Objective feedback derived from metrics like player load and total impacts supports both performance reviews and recovery protocols. Quantitative insights align player perceptions of exertion with actual workload, guiding nutrition, rest and physiotherapy interventions for efficient return to peak readiness.
    """
    )

    doc.add_paragraph("Speed Zones")

    # Get dynamic speed zones from config
    zones_dict = get_speed_zones(season)
    
    zones_data = {
        "Zone": list(zones_dict.keys()),
        "Threshold (km/h)": [v["range"] for v in zones_dict.values()],
        "Intensity": [v["intensity"] for v in zones_dict.values()],
        "Descriptor": [v["descriptor"] for v in zones_dict.values()],
    }

    zones_df = pd.DataFrame(zones_data)
    add_dataframe_as_table(doc, zones_df)


def add_challenges_section(doc: Document) -> None:
    """
    Add standardized Challenges section to report.

    Lists common operational and data management issues.

    Args:
        doc: python-docx Document object
    """
    doc.add_heading("Challenges", level=1)
    doc.add_paragraph(
        """The challenges outlined below represent common operational and data management issues observed across multiple clubs using GPS tracking systems. These are not specific to any one club, but rather reflect broader trends that can impact the consistency, accuracy, and usefulness of performance data. From equipment mismanagement to gaps in staff training and session documentation, these hurdles can hinder the full potential of tracking insights. Clubs experiencing similar difficulties or seeking tailored support are encouraged to contact the performance analytics team, who can provide club-specific guidance and help implement effective solutions.

The challenges are outlined below:

1. Inconsistent use of tracking devices across games, leading to incomplete match coverage
2. Failure to upload sessions at all or delays in syncing which creates backlogs
3. Misclassification of training sessions as games, and vice versa
4. Incorrect session naming, making it difficult to locate or aggregate sessions
5. Improper session splitting (e.g., durations recorded as < 90 minutes or > 120 minutes, incorrect splits for substituted players)
6. Mismanagement of equipment leading to loss or damage
7. Limited expertise in translating raw GPS outputs into tactical or physiological insights
8. Club staff not fully trained on usage and upload procedures, resulting in missed or late uploads
        """
    )


def add_future_plans_section(doc: Document, season: str = "2025/26") -> None:
    """
    Add standardized Future Plans section to report.

    Outlines federation-level initiatives.

    Args:
        doc: python-docx Document object
        season: Football season (e.g., '2025/26')
    """
    doc.add_heading("Future Plans", level=1)
    doc.add_paragraph(
        """This section outlines strategic initiatives to maximize the impact of GPS tracking data across the top-tier leagues. These actions build on the {} season insights and are designed to standardize workflows, deepen analysis, and support both club staff and players. For club-specific implementation advice or technical assistance, please contact the performance analytics team.
""".format(season) + """
1. Conduct regular visits to the clubs to monitor equipment usage and share club specific insights.
2. Organise regular training workshops for club performance staff on GPS device management, data procedures and software analytics to ensure consistent, high-quality data capture.
3. Leverage tracking data to inform nutritional guidance, tailoring macronutrient ratios to individual training demands and recovery needs.
4. Introduce individual performance benchmarks for every playing position (right back, centre forward, defensive midfielder, etc.) with detailed threshold bands for both volume and intensity at club level.
5. Integrate GPS metrics with tactical and technical analysis by linking tracking outputs to technical or tactical data generated by other systems, thereby providing coaches with combined insights into player movement patterns and team strategy.
6. Carry out research projects using tracking data along other data sources to improve planning.
        """
    )


def add_conclusion_section(doc: Document, season: str = "2025/26") -> None:
    """
    Add standardized Conclusion section to report.

    Thanks stakeholders, references data privacy policy, provides contact info.

    Args:
        doc: python-docx Document object
        season: Football season (e.g., '2025/26')
    """
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(
        """We extend our sincere thanks to every coach, performance staff member and analyst for their dedication in collecting and interpreting GPS tracking data during the {} season. Your hard work has laid a solid foundation for evidence-based decision-making across fitness, recovery and match preparation.
    """.format(season) + """
All tracking outputs and associated player information will be handled in strict accordance with the data privacy and security policy, ensuring confidentiality, ethical use and compliant storage. This protocol safeguards both individual rights and the integrity of our performance analysis.

As we move into the next season, we look forward to your feedback and stand ready to support your club through the performance analytics team. You can reach us via email at analytics@example.com
        """
    )


def embed_matplotlib_figure(doc: Document, fig, width_inches: float = 6.0, caption: str = None) -> None:
    """
    Embed a matplotlib figure in the document with optional numbered caption.

    Args:
        doc: python-docx Document object
        fig: matplotlib figure object
        width_inches: Width of embedded image (default: 6.0)
        caption: Optional caption text
    """
    img_stream = BytesIO()
    try:
        # Check if figure is empty (no axes)
        if hasattr(fig, 'axes') and not fig.axes:
            # Create a dummy figure with text
            import matplotlib.pyplot as plt
            dummy_fig, ax = plt.subplots(figsize=(6, 2))
            ax.text(0.5, 0.5, "No data available for visualization", 
                   ha='center', va='center', fontsize=12)
            ax.axis('off')
            dummy_fig.savefig(img_stream, format="png", bbox_inches="tight", dpi=300)
            plt.close(dummy_fig)
        else:
            fig.savefig(img_stream, format="png", bbox_inches="tight", dpi=300)
            
        img_stream.seek(0)
        
        # Add to document with centering
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(img_stream, width=Inches(width_inches))
        
        if caption:
            if not hasattr(doc, '_figure_count'):
                doc._figure_count = 0
            doc._figure_count += 1
            
            paragraph = doc.add_paragraph()
            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = paragraph.add_run(f"Figure {doc._figure_count}: {caption}")
            run.italic = True
            run.font.size = doc.styles['Normal'].font.size

    except Exception as e:
        print(f"Error embedding figure: {e}")
        # Insert a placeholder text in doc instead
        doc.add_paragraph(f"[Visualisation failed: {e}]")


def embed_matplotlib_axis(doc: Document, ax, width_inches: float = 6.0) -> None:
    """
    Embed a matplotlib axes object in the document.

    Args:
        doc: python-docx Document object
        ax: matplotlib axes object
        width_inches: Width of embedded image (default: 6.0)
    """
    img_stream = BytesIO()
    ax.figure.savefig(img_stream, format="png", bbox_inches="tight", dpi=100)
    img_stream.seek(0)
    doc.add_picture(img_stream, width=Inches(width_inches))


def save_document(doc: Document, output_dir: str, filename: str) -> str:
    """
    Save document to specified location.

    Creates output directory if it doesn't exist.

    Args:
        doc: python-docx Document object
        output_dir: Directory path for output
        filename: Filename (e.g., 'report.docx')

    Returns:
        Full path to saved document
    """
    os.makedirs(output_dir, exist_ok=True)
    output_path = os.path.join(output_dir, filename)
    doc.save(output_path)
    return output_path


def create_report_document(club_name: str) -> Document:
    """
    Create a new report document with title page.

    Args:
        club_name: Name of club for report title

    Returns:
        python-docx Document object
    """
    doc = Document()
    doc.add_heading(f"{club_name} Performance Report", 0)
    doc.add_page_break()
    return doc
