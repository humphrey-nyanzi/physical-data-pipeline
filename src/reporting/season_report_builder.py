"""
Season Report Builder Module

Orchestrates the creation of the Season/Half-Season report.
Uses analysis statistics and visualizations to populate a DOCX template.
"""

import logging
import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH # Unused, but leaving if needed for table alignment later? Actually Ruff says unused.
# Removing it:


from ..analysis import season_analysis
from ..analysis import visualizations
from . import document_generation as doc_gen
from . import report_builder
from ..config import constants
from src.config.styles import ReportStyles
import yaml
from pathlib import Path

logger = logging.getLogger(__name__)

class SeasonReportBuilder:
    def __init__(
        self, 
        df: pd.DataFrame, 
        league: str, 
        timeframe: str, 
        output_dir: Path,
        half_season_md: int,
        season: str = "2025/26",
        gk_mode: bool = False,
        run_id: str = "manual"
    ):
        self.df = df
        self.league = league
        self.timeframe = timeframe
        self.output_dir = output_dir
        self.half_season_md = half_season_md
        self.season = season
        self.gk_mode = gk_mode
        self.run_id = run_id
        
        # Load config
        self.config_path = Path(__file__).parents[1] / "config" / "analysis_config.yaml"
        
        # Safe loading
        if self.config_path.exists():
            with open(self.config_path, 'r') as f:
                self.config = yaml.safe_load(f)
        else:
             self.config = {} # improved error handling during refactor

        self.doc = Document()
        self._configure_styles()
    
    def _configure_styles(self):
        """Apply centralized styles."""
        ReportStyles.apply_normal_style(self.doc)
        ReportStyles.apply_heading_styles(self.doc)
        
    def build_report(self, suffix: str = ""):
        """Execute the full report generation pipeline."""
        logger.info(f"Building {self.timeframe} report for {self.league} (suffix: {suffix})...")
        
        # 1. Title Page
        assets_dir = Path(__file__).parent / "assets"
        base_title = f"{self.league.upper()} {self.timeframe.replace('_', ' ').title()} Report"
        display_title = f"Goalkeeper {base_title}" if self.gk_mode else base_title
        
        doc_gen.add_title_page(
            self.doc, 
            title=display_title,
            subtitle=f"{self.league} | {self.season} Season",
            logo_path=str(assets_dir / "fufa_logo.png")
        )
        
        # Insert Table of Contents
        doc_gen.add_table_of_contents(self.doc)
        
        # 2. Introduction
        doc_gen.add_branded_header(self.doc, "Introduction", icon_path=None)
        self._add_intro_section()
        
        # 3. Usage Analysis
        self._add_usage_section()
        
        # 4. Performance Analysis
        self._add_performance_section()
        
        # 5. Conclusion
        self._add_conclusion_section()
        
        # Save
        # Standardized naming: {League}_{Season}_{Scope}_{Summary}_{Run_ID}.docx
        season_filename = self.season.replace("/", "-")
        scope = self.timeframe.title().replace('_', '')
        filename = f"{self.league.upper()}_{season_filename}_{scope}_Report_{self.run_id}{suffix}.docx"
        output_path = self.output_dir / filename
        self.doc.save(output_path)
        logger.info(f"Report saved to {output_path}")
        return output_path


    def _add_intro_section(self):
        self.doc.add_paragraph(
            f"This report presents a comprehensive analysis of the {self.league.upper()} "
            f"competitions for the {self.season} {self.timeframe.replace('_', ' ')}. "
            "The data collected via Catapult wearable technology provides insights into the physical demands "
            "and performance characteristics of the league."
        )
        self.doc.add_paragraph(
            "The report is divided into two main sections: Usage Analysis, which details data coverage "
            "and participation, and Performance Analysis, which delves into the physical metrics."
        )

    # ========================================================================
    # USAGE SECTION
    # ========================================================================
    def _add_usage_section(self):
        doc_gen.add_branded_header(self.doc, "Usage Analysis", icon_path=None)
        self.doc.add_paragraph(
            "This section outlines the extent of data collection, highlighting club participation "
            "and consistency across matchdays."
        )
        
        usage_stats = season_analysis.get_usage_stats(self.df)
        
        # 2.1 Unique Players per Club
        self.doc.add_heading("2.1 Player Participation per Club", level=2)
        unique_players = usage_stats['unique_players_per_club'].rename(columns=constants.GENERAL_DISPLAY_NAMES)
        top_club = unique_players.iloc[0]
        doc_gen.add_dataframe_as_table(self.doc, unique_players, caption="Unique Players Analyzed per Club")
        self.doc.add_paragraph(
            f"{top_club['Club']} had the highest number({top_club['Number of Players']}) of unique players analyzed, "
            "indicating significant squad rotation or broad data capture."
        )
        
        # 2.2 Participation Trends (Players & Clubs per Matchday)
        self.doc.add_heading("2.2 Matchday Participation Trends", level=2)
        
        # Plot Players per MD
        fig = visualizations.plot_league_trend(
            usage_stats['players_per_md_trend'], 'total_players', 'Total Unique Players Captured'
        )
        doc_gen.embed_matplotlib_figure(self.doc, fig, caption="Total Unique Players Captured per Matchday")
        plt.close(fig)
        
        # Plot Clubs per MD
        fig = visualizations.plot_league_trend(
            usage_stats['clubs_per_md_trend'], 'total_clubs', 'Total Clubs Submitting Data'
        )
        doc_gen.embed_matplotlib_figure(self.doc, fig, caption="Total Clubs Submitting Data per Matchday")
        plt.close(fig)
        
        self.doc.add_paragraph(
            "The charts above illustrate the volume of data captured over time. "
            "Peaks typically correspond to full matchday rounds where most clubs uploaded data."
        )
        
        # 2.3 Coverage Heatmap
        self.doc.add_heading("2.3 Data Coverage Grid", level=2)
        self.doc.add_paragraph(
            "The grid below visualizes which clubs uploaded data for each matchday (Green = Uploaded)."
        )
        grid_df = season_analysis.get_coverage_grid(self.df)
        fig = visualizations.plot_coverage_heatmap(grid_df)
        doc_gen.embed_matplotlib_figure(self.doc, fig, caption="Matchday Data Coverage Grid (Green = Data Available)")
        plt.close(fig)

    # ========================================================================
    # PERFORMANCE SECTION
    # ========================================================================
    def _add_performance_section(self):
        doc_gen.add_branded_header(self.doc, "Performance Analysis", icon_path=None)
        self.doc.add_paragraph(
            "This section analyzes the physical output of players, categorized by volume "
            "(total work done) and intensity (rate of work)."
        )
        
        perf_stats = season_analysis.get_performance_stats(self.df)
        
        # 3.1 Metric Distributions (Histograms)
        self.doc.add_heading("3.1 Metric Distributions", level=2)
        self.doc.add_paragraph("Distribution of key volume (Distance) and intensity (Player Load) metrics.")
        
        # Distance Histogram
        fig = visualizations.plot_metric_histogram(self.df, 'distance_km', 'Total Distance (km)')
        doc_gen.embed_matplotlib_figure(self.doc, fig, caption="Distribution of Total Distance (km) across all sessions")
        plt.close(fig)
        
        # Player Load Histogram
        fig = visualizations.plot_metric_histogram(self.df, 'player_load', 'Player Load')
        doc_gen.embed_matplotlib_figure(self.doc, fig, caption="Distribution of Player Load across all sessions")
        plt.close(fig)
        
        # 3.2 Aggregated Metrics
        self.doc.add_heading("3.2 League Averages", level=2)
        self.doc.add_heading("Volume Metrics", level=3)
        vol_avg = perf_stats['volume'].copy()
        vol_avg['Metric'] = vol_avg['Metric'].map(constants.METRIC_DISPLAY_NAMES).fillna(vol_avg['Metric'])
        doc_gen.add_dataframe_as_table(self.doc, vol_avg.round(2), caption="League-wide Volume Metrics Averages")
        
        self.doc.add_heading("Intensity Metrics", level=3)
        int_avg = perf_stats['intensity'].copy()
        int_avg['Metric'] = int_avg['Metric'].map(constants.METRIC_DISPLAY_NAMES).fillna(int_avg['Metric'])
        doc_gen.add_dataframe_as_table(self.doc, int_avg.round(2), caption="League-wide Intensity Metrics Averages")
        
        # 3.3 Positional Breakdown
        self.doc.add_heading("3.3 League Averages by Position", level=2)
        pos_stats = report_builder.get_average_metrics_by_position(
            self.df, 
            constants.VOLUME_METRICS, 
            constants.INTENSITY_METRICS, 
            constants.METRIC_DISPLAY_NAMES
        ).rename(columns=constants.GENERAL_DISPLAY_NAMES)
        doc_gen.add_dataframe_as_table(self.doc, pos_stats.round(2), caption="League-wide Averages by Playing Position")
        
        # 3.4 Club Level Comparisons
        self.doc.add_heading("3.4 Club Level Comparisons", level=2)
        club_stats = season_analysis.get_club_comparison(self.df)
        
        if 'volume' in club_stats:
            self.doc.add_heading("Volume Metrics by Club", level=3)
            vol_df = club_stats['volume'].rename(columns=constants.METRIC_DISPLAY_NAMES).reset_index().rename(columns=constants.GENERAL_DISPLAY_NAMES)
            doc_gen.add_dataframe_as_table(self.doc, vol_df.round(2), caption="Volume Metrics Averages by Club")
            
        if 'intensity' in club_stats:
            self.doc.add_heading("Intensity Metrics by Club", level=3)
            int_df = club_stats['intensity'].rename(columns=constants.METRIC_DISPLAY_NAMES).reset_index().rename(columns=constants.GENERAL_DISPLAY_NAMES)
            doc_gen.add_dataframe_as_table(self.doc, int_df.round(2), caption="Intensity Metrics Averages by Club")
        
        # 3.5 Contextual Analysis
        self.doc.add_heading("3.5 Contextual Analysis", level=2)
        context_stats = season_analysis.get_contextual_stats(self.df)
        
        if 'location' in context_stats:
            self.doc.add_heading("Home vs Away", level=3)
            # Plot specific metric comparison, e.g. PL
            metric_label = constants.METRIC_DISPLAY_NAMES.get('player_load', 'Player Load')
            fig = visualizations.plot_context_comparison(
                context_stats['location'], 'player_load', metric_label, 'Location'
            )
            doc_gen.embed_matplotlib_figure(self.doc, fig, caption=f"Comparison of {metric_label} (Home vs Away)")
            plt.close(fig)
            
        if 'result' in context_stats:
            self.doc.add_heading("Match Result (Win/Draw/Loss)", level=3)
            metric_label = constants.METRIC_DISPLAY_NAMES.get('distance_km', 'Distance (km)')
            fig = visualizations.plot_context_comparison(
                context_stats['result'], 'distance_km', metric_label, 'Result'
            )
            doc_gen.embed_matplotlib_figure(self.doc, fig, caption=f"Comparison of {metric_label} covered by Match Result")
            plt.close(fig)
            
        # 3.6 Rolling Trends
        self.doc.add_heading("3.6 Seasonal Metric Trends", level=2)
        self.doc.add_paragraph("Rolling averages (3-matchday window) for key performance indicators.")
        metrics_map = {
            'distance_km': 'Distance (km)',
            'player_load': 'Player Load',
            'top_speed_kmh': 'Top Speed (km/h)',
            'work_ratio': 'Work Ratio'
        }
        doc_gen.set_landscape(self.doc)
        fig = visualizations.plot_rolling_trend_grid(self.df, metrics_map)
        doc_gen.embed_matplotlib_figure(self.doc, fig, width_inches=9.0, caption="Seasonal Performance Trends (3-Match Rolling Average)")
        plt.close(fig)
        doc_gen.set_portrait(self.doc)
        
        # 3.7 Speed Zones
        self.doc.add_heading("3.7 Speed Zone Analysis", level=2)
        sz_stats = season_analysis.get_speed_zone_stats(self.df)
        if 'dist_pct' in sz_stats:
            self.doc.add_heading("Distance in Speed Zones", level=3)
            fig = visualizations.plot_speed_zones_stacked(
                sz_stats['dist_pct'], "Distance Distribution by Speed Zone"
            )
            doc_gen.embed_matplotlib_figure(self.doc, fig, caption="Distance Distribution by Speed Zone across the League")
            plt.close(fig)
            
        self._add_record_breakers()
        
    def _add_record_breakers(self):
        self.doc.add_heading("3.8 Record Breakers", level=2)
        self.doc.add_paragraph(
            "The following tables highlight the highest single-match values recorded for key metrics."
        )
        
        metrics = ['distance_km', 'top_speed_kmh', 'player_load', 'sprint_distance_m']
        records = season_analysis.get_max_performers(self.df, metrics)
        
        for metric, df in records.items():
            if not df.empty:
                display_metric = constants.METRIC_DISPLAY_NAMES.get(metric, metric)
                top_p = df.iloc[0]
                self.doc.add_paragraph(
                    f"Highest {display_metric}: {top_p['p_name']} ({top_p['club_for']}) - {top_p[metric]:.2f}"
                )
                safe_df = df.copy().rename(columns={
                    **constants.GENERAL_DISPLAY_NAMES,
                    metric: display_metric
                })
                doc_gen.add_dataframe_as_table(self.doc, safe_df.round(2), caption=f"Top 5 Performers for {display_metric}")

    def _add_conclusion_section(self):
        doc_gen.add_conclusion_section(self.doc, season=self.season)
