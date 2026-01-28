import os
from datetime import datetime
from docx import Document
from src.config.styles import ReportStyles
from .document_generation import add_table_of_contents
import pandas as pd

class WeeklyGPSReportBuilder:
    def __init__(self, matchday_number, season="2025/2026", gk_mode=False, league="upl"):
        self.matchday_number = matchday_number
        self.season = season
        self.gk_mode = gk_mode
        self.league = league
        self.doc = Document()
        self._configure_styles()
        
        # Configuration
        self.TOP_N = 5

    def _configure_styles(self):
        """Configure document styles using centralized configuration."""
        ReportStyles.apply_normal_style(self.doc)
        ReportStyles.apply_heading_styles(self.doc)

    def build_report(self, df_all, uploading_teams, missing_teams):
        """Build the complete report."""
        self._add_title()
        # add_table_of_contents(self.doc) Not necessary for a weekly report
        self._add_introduction() 
        self._add_top_performers(df_all)
        self._add_matchday_averages(df_all, uploading_teams)
        self._add_missing_teams(missing_teams)

    def _add_title(self):
        report_type = "GOALKEEPER" if self.gk_mode else "PHYSICAL PERFORMANCE (CATAPULT)"
        league_name = "UGANDA PREMIER LEAGUE" if self.league.lower() == "upl" else "FUFA WOMEN SUPER LEAGUE"
        title_text = f"GPS {report_type} REPORT FOR {league_name} {self.season} MATCHDAY {self.matchday_number}"
        self.doc.add_paragraph(title_text, style='Title')
        self.doc.add_page_break()

    def _add_introduction(self):
        self.doc.add_heading("1. Introduction", level=1)
        league_full_name = "Uganda Premier League" if self.league.lower() == "upl" else "FUFA Women's Super League"
        self.doc.add_paragraph(
            f"This report provides a detailed analysis of the physical performance data captured via Catapult GPS technology "
            f"during Matchday {self.matchday_number} of the {league_full_name} {self.season} season."
        )
        self.doc.add_paragraph(
            "The objective of this weekly report is to highlight top performers across key physical metrics, "
            "provide matchday averages for benchmarking, and track data submission compliance across all league teams."
        )

    def _add_top_performers(self, df_all):
        self.doc.add_heading(f"2. TOP {self.TOP_N} PERFORMERS PER METRIC", level=1)
        
        # Disclaimer
        disclaimer = self.doc.add_paragraph()
        runner = disclaimer.add_run("NOTE: Sprint distances should be viewed with caution as speed zones are still being adjusted for various clubs.")
        runner.bold = True
        runner.font.color.rgb = ReportStyles.COLOR_RED
        
        metrics = {
            'Distance (km)': ('Total Distance', 'km', 2),
            'Sprint Distance (m)': ('Sprint Distance', 'm', 1),
            'Top Speed (km/h)': ('Top Speed', 'km/h', 2),
            'Player Load': ('Player Load', '', 1)
        }
        
        for col, (label, unit, decimals) in metrics.items():
            self.doc.add_heading(label, level=2)
            
            if col not in df_all.columns:
                self.doc.add_paragraph("Metric data not found")
                continue

            top = df_all.nlargest(self.TOP_N, col)
            
            if top.empty:
                self.doc.add_paragraph("No data available")
                continue
            
            # Prepare display DataFrame
            display_rows = []
            for i, (idx, row) in enumerate(top.iterrows(), 1):
                raw_clean_name = row.get('Clean Name', row.get('Player Name', ''))
                display_rows.append({
                    "S/N": i,
                    "Player Name": str(raw_clean_name).title(),
                    "Club": row.get('team1', ''),
                    "Position": row.get('Position', ''),
                    f"Value ({unit})" if unit else "Value": row[col]
                })
            
            top_df = pd.DataFrame(display_rows)
            
            # Add Table using unified generator
            from . import document_generation as doc_gen
            doc_gen.add_dataframe_as_table(self.doc, top_df)

    def _add_matchday_averages(self, df_all, uploading_teams):
        self.doc.add_page_break()
        self.doc.add_heading("3. Matchday Averages (Exposure Filtered)", level=1)
        
        from src.config import constants
        min_dur = constants.MIN_SESSION_DURATION_MINUTES
        min_dist = constants.MIN_SESSION_DISTANCE_KM
        
        # df_all is already filtered by the pipeline
        df_filtered = df_all
        
        self.doc.add_paragraph(
            f"Filter: Duration ≥ {min_dur} min AND Distance ≥ {min_dist} km"
        )
        self.doc.add_paragraph(f"Players included: {len(df_filtered)}")

        self.doc.add_paragraph(f"Report Generated on: {datetime.now():%Y-%m-%d %H:%M}")
        self.doc.add_paragraph(f"Teams Analysed: {len(uploading_teams)}")
        
        self.doc.add_paragraph("")
        
        metrics = {
            'Distance (km)': ('Total Distance', 'km', 2),
            'Sprint Distance (m)': ('Sprint Distance', 'm', 1),
            'Top Speed (km/h)': ('Top Speed', 'km/h', 2),
            'Player Load': ('Player Load', '', 1)
        }

        if df_filtered.empty:
            self.doc.add_paragraph("⚠ No players met exposure criteria")
        else:
            for col, (label, unit, decimals) in metrics.items():
                if col not in df_filtered.columns:
                    continue
                    
                mean_val = df_filtered[col].mean()
                median_val = df_filtered[col].median()
                std_val = df_filtered[col].std()
                count = df_filtered[col].count()
                
                if unit:
                    text = f"{label}: Mean = {mean_val:.{decimals}f} {unit}, Median = {median_val:.{decimals}f} {unit}, SD = {std_val:.{decimals}f}, n = {count}"
                else:
                    text = f"{label}: Mean = {mean_val:.{decimals}f}, Median = {median_val:.{decimals}f}, SD = {std_val:.{decimals}f}, n = {count}"
                
                self.doc.add_paragraph(text, style='List Bullet')

    def _add_missing_teams(self, missing_teams):
        if missing_teams:
            self.doc.add_heading("4. Teams Not Uploading Data", level=1)
            self.doc.add_paragraph("The following teams did not submit their (Catapult) GPS data:")
            for team in sorted(list(missing_teams)):
                self.doc.add_paragraph(f"{team}", style='List Bullet')

    def save(self, output_path):
        """Save the document."""
        # Ensure directory exists
        os.makedirs(os.path.dirname(output_path), exist_ok=True)
        self.doc.save(output_path)
